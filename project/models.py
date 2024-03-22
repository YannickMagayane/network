from __future__ import unicode_literals

from django.core.mail import send_mail
from django.db import models
from django.contrib.auth.base_user import AbstractBaseUser
from django.contrib.auth.models import PermissionsMixin
from django.utils.translation import gettext_lazy as _
from django.contrib.auth.base_user import BaseUserManager
from scapy.all import sniff, IP, TCP
from django.db.models.signals import post_save
from django.dispatch import receiver
from PyPDF2 import PdfFileReader
from docx import Document
from PIL import Image
from io import BytesIO
import re


class UserManager(BaseUserManager):
    use_in_migrations = True

    def _create_user(self, email, password, **extra_fields):
        if not email:
            raise ValueError('Vous devez mettre une adresse mail')
        email = self.normalize_email(email)
        user = self.model(email=email, **extra_fields)
        user.set_password(password)
        user.save(using=self._db)
        return user

    def create_user(self, email, password=None, **extra_fields):
        extra_fields.setdefault('is_superuser', False)
        return self._create_user(email, password, **extra_fields)

    def create_superuser(self, email, password, **extra_fields):
        extra_fields.setdefault('is_superuser', True)

        if extra_fields.get('is_superuser') is not True:
            raise ValueError('Le superuser n\'est pas activé')

        return self._create_user(email, password, **extra_fields)


class User(AbstractBaseUser, PermissionsMixin):

    email = models.EmailField(_('email address'), unique=True)
    first_name = models.CharField(_('first name'), max_length=30, blank=True)
    last_name = models.CharField(_('last name'), max_length=30, blank=True)
    date_joined = models.DateTimeField(_('date joined'), auto_now_add=True)
    is_active = models.BooleanField(_('active'), default=True)
    is_staff = models.BooleanField(_('staff'), default=True)

    objects = UserManager()

    USERNAME_FIELD = 'email'
    REQUIRED_FIELDS = []

    class Meta:
        verbose_name = _('user')
        verbose_name_plural = _('users')

    def get_full_name(self, full_name=None):
        nom_complet = '%s %s' % (self.first_name, self.last_name)
        return nom_complet.strip()

    def get_short_name(self):
        return self.first_name

    def email_user(self, subject, message, from_email=None, **kwargs):
        send_mail(subject, message, from_email, [self.email], **kwargs)












class Fichier(models.Model):
    nom_du_fichier = models.CharField(max_length=50)
    fichier = models.FileField(upload_to='fichiers/')  # Champ pour stocker le fichier envoyé
    user = models.ForeignKey(User, on_delete=models.CASCADE)

    def __str__(self):
        return self.nom_du_fichier




class IntrusionDetectionLog(models.Model):
    timestamp = models.DateTimeField(auto_now_add=True)
    source_ip = models.CharField(max_length=50)
    destination_ip = models.CharField(max_length=50)
    intrusion_type = models.CharField(max_length=100, blank=True, null=True)
    description = models.TextField(blank=True, null=True)
    is_malicious = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.timestamp} - {self.intrusion_type}"

    @staticmethod
    def detect_intrusion(packet):
        if IP in packet and TCP in packet:
            if packet[TCP].flags == 2:  # SYN Flag
                return "SYN Flood", "Potential SYN Flood detected"
            # Ajoutez d'autres logiques de détection d'intrusion ici selon vos besoins
        return None, None

    @classmethod
    def capture_packets(cls):
        packets = sniff(count=100)
        for packet in packets:
            intrusion_type, description = cls.detect_intrusion(packet)
            if intrusion_type and description:
                cls.objects.create(
                    source_ip=packet[IP].src,
                    destination_ip=packet[IP].dst,
                    intrusion_type=intrusion_type,
                    description=description,
                    is_malicious=True  # Marquer comme malveillant automatiquement
                )

    @staticmethod
    def is_pdf(file):
        try:
            PdfFileReader(BytesIO(file)).getNumPages()
            return True
        except Exception as e:
            return False

    @staticmethod
    def is_docx(file):
        try:
            Document(BytesIO(file))
            return True
        except Exception as e:
            return False

    @staticmethod
    def is_png(file):
        try:
            Image.open(BytesIO(file))
            return True
        except Exception as e:
            return False

    @classmethod
    def detect_file_intrusion(cls, file):
        if cls.is_pdf(file):
            if cls.has_links_pdf(file):
                return "PDF Intrusion", "Potential PDF Intrusion detected"
        elif cls.is_docx(file):
            if cls.has_links_docx(file):
                return "DOCX Intrusion", "Potential DOCX Intrusion detected"
        elif cls.is_png(file):
            if cls.has_links_png(file):
                return "PNG Intrusion", "Potential PNG Intrusion detected"
        return None, None

    @staticmethod
    def has_links_pdf(file):
        try:
            reader = PdfFileReader(BytesIO(file))
            for page in reader.pages:
                if re.search(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', page.extract_text()):
                    return True
            return False
        except Exception as e:
            return False

    @staticmethod
    def has_links_docx(file):
        try:
            doc = Document(BytesIO(file))
            for paragraph in doc.paragraphs:
                if re.search(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', paragraph.text):
                    return True
            return False
        except Exception as e:
            return False

    @staticmethod
    def has_links_png(file):
        # Vous pouvez adapter cette méthode en fonction de vos besoins spécifiques pour les images PNG
        return False

# Déclenche la capture de paquets lors de la création d'un nouveau log
models.signals.post_save.connect(IntrusionDetectionLog.capture_packets, sender=IntrusionDetectionLog)